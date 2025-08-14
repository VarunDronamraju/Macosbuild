import os
import uuid
from pathlib import Path
from typing import List, Dict, Optional
import PyPDF2
from docx import Document as DocxDocument
from sentence_transformers import SentenceTransformer
import numpy as np
from backend.config import settings
from backend.database import SessionLocal, Document, DocumentChunk, User

class DocumentProcessor:
    def __init__(self):
        self.embedding_model = SentenceTransformer(settings.EMBEDDING_MODEL)
    
    def extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text from various file formats"""
        try:
            if file_type.lower() == 'pdf':
                return self._extract_pdf_text(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return self._extract_docx_text(file_path)
            elif file_type.lower() == 'txt':
                return self._extract_txt_text(file_path)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Text extraction failed: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text.strip()
    
    def _extract_txt_text(self, file_path: str) -> str:
        """Extract text from TXT"""
        # Try different encodings to handle various text files
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
        # If all encodings fail, try with error handling
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
            return file.read().strip()
    
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into chunks with overlap"""
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings near the chunk boundary
                for i in range(end, max(start + chunk_size // 2, end - 100), -1):
                    if text[i] in '.!?':
                        end = i + 1
                        break
            
            chunk = text[start:end].strip()
            if len(chunk) > 50:  # Filter out very short chunks
                chunks.append(chunk)
            
            if end >= len(text):
                break
                
            start = end - overlap
        
        return chunks
    
    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """Generate embeddings for text chunks"""
        return self.embedding_model.encode(texts)
    
    def process_document(self, file_path: str, filename: str, user_id: str) -> str:
        """Complete document processing pipeline"""
        db = SessionLocal()
        try:
            # CRITICAL FIX: Handle user creation for test users or non-existent users
            if user_id == "test_user":
                # For test_user, use SAME string "test_user" for BOTH database AND vector storage
                test_user_uuid = uuid.UUID("00000000-0000-0000-0000-000000000001")
                existing_user = db.query(User).filter(User.id == test_user_uuid).first()
                if not existing_user:
                    test_user = User(
                        id=test_user_uuid,
                        email="test@example.com",
                        name="Test User"
                    )
                    db.add(test_user)
                    db.commit()
                    db.refresh(test_user)
                
                # Use UUID for database, but keep "test_user" string for vector collection
                db_user_id = str(test_user_uuid)
                vector_user_id = "test_user"  # Keep as string for vector operations
                print(f"DEBUG DOCS: Using db_user_id={db_user_id}, vector_user_id={vector_user_id}")
            else:
                # For real users, use UUID for both
                try:
                    user_uuid = uuid.UUID(user_id)
                    existing_user = db.query(User).filter(User.id == user_uuid).first()
                    if not existing_user:
                        new_user = User(
                            id=user_uuid,
                            email=f"user_{user_id[:8]}@example.com",
                            name=f"User {user_id[:8]}"
                        )
                        db.add(new_user)
                        db.commit()
                        db.refresh(new_user)
                    db_user_id = str(user_uuid)
                    vector_user_id = str(user_uuid)
                except ValueError:
                    # Invalid UUID format, create a new user
                    user_uuid = uuid.uuid4()
                    new_user = User(
                        id=user_uuid,
                        email=f"user_{user_id[:8]}@example.com",
                        name=f"User {user_id[:8]}"
                    )
                    db.add(new_user)
                    db.commit()
                    db.refresh(new_user)
                    db_user_id = str(user_uuid)
                    vector_user_id = str(user_uuid)

            # Create document record
            file_type = Path(filename).suffix[1:].lower()
            file_size = os.path.getsize(file_path)
            
            document = Document(
                filename=filename,
                file_type=file_type,
                file_size=file_size,
                user_id=db_user_id,  # Use UUID for database
                processing_status="processing"
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            
            print(f"DEBUG DOCS: Created document {document.id} for user {db_user_id}")
            
            # Extract text
            text = self.extract_text(file_path, file_type)
            print(f"DEBUG DOCS: Extracted {len(text)} characters from {filename}")
            
            # Chunk text
            chunks = self.chunk_text(text)
            print(f"DEBUG DOCS: Created {len(chunks)} chunks")
            
            # Generate embeddings
            embeddings = self.generate_embeddings(chunks)
            print(f"DEBUG DOCS: Generated embeddings shape: {embeddings.shape}")
            
            # Store chunks and embeddings
            from backend.rag import VectorStore
            vector_store = VectorStore()
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                embedding_id = str(uuid.uuid4())
                
                print(f"DEBUG DOCS: Storing chunk {i} with ID {embedding_id}")
                print(f"DEBUG DOCS: Chunk preview: {chunk[:100]}...")
                print(f"DEBUG DOCS: Using vector_user_id: {vector_user_id}")
                
                # CRITICAL: Store in vector database using the SAME user_id format as search
                vector_store.store_embedding(
                    embedding_id, 
                    embedding, 
                    {
                        "document_id": str(document.id),
                        "chunk_index": i, 
                        "content": chunk[:200] + "..." if len(chunk) > 200 else chunk
                    },
                    vector_user_id  # Use consistent user_id for vector storage
                )
                
                # Store chunk in SQL database
                chunk_record = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk,
                    embedding_id=embedding_id
                )
                db.add(chunk_record)
                print(f"DEBUG DOCS: Stored chunk {i} in database")
            
            # Update document status
            document.processing_status = "completed"
            document.chunk_count = len(chunks)
            db.commit()
            
            print(f"DEBUG DOCS: Document processing completed with {len(chunks)} chunks")
            return str(document.id)
            
        except Exception as e:
            print(f"ERROR DOCS: {e}")
            # Rollback the session to handle any pending transactions
            db.rollback()
            
            # Try to update document status if it was created
            try:
                if 'document' in locals():
                    # Get a fresh session for the status update
                    status_db = SessionLocal()
                    try:
                        doc = status_db.query(Document).filter(Document.id == document.id).first()
                        if doc:
                            doc.processing_status = "failed"
                            status_db.commit()
                    except:
                        status_db.rollback()
                    finally:
                        status_db.close()
            except:
                pass  # Ignore errors in status update
                
            raise e
        finally:
            db.close()

# Utility functions
def save_uploaded_file(file_content: bytes, filename: str) -> str:
    """Save uploaded file to disk"""
    file_path = settings.UPLOAD_DIR / filename
    with open(file_path, 'wb') as f:
        f.write(file_content)
    return str(file_path)

def get_user_documents(user_id: str) -> List[Dict]:
    """Get all documents for a user"""
    db = SessionLocal()
    try:
        # Convert user_id to UUID if it's a valid UUID string
        try:
            user_uuid = uuid.UUID(user_id)
        except ValueError:
            # If it's not a valid UUID (like "test_user"), use the test user ID
            if user_id == "test_user":
                user_uuid = uuid.UUID("00000000-0000-0000-0000-000000000001")
            else:
                # Return empty list for invalid user IDs
                return []
        
        documents = db.query(Document).filter(Document.user_id == user_uuid).all()
        return [
            {
                "id": str(doc.id),
                "filename": doc.filename,
                "file_type": doc.file_type,
                "file_size": doc.file_size,
                "upload_date": doc.upload_date.isoformat(),
                "processing_status": doc.processing_status,
                "chunk_count": doc.chunk_count
            }
            for doc in documents
        ]
    except Exception as e:
        print(f"Error in get_user_documents: {e}")
        return []
    finally:
        db.close()

def delete_user_document(document_id: str, user_id: str) -> bool:
    """Delete a document and its chunks"""
    db = SessionLocal()
    try:
        # Convert user_id if needed
        try:
            user_uuid = uuid.UUID(user_id)
        except ValueError:
            if user_id == "test_user":
                user_uuid = uuid.UUID("00000000-0000-0000-0000-000000000001")
            else:
                return False

        # Find document
        document = db.query(Document).filter(
            Document.id == document_id,
            Document.user_id == user_uuid
        ).first()
        
        if not document:
            return False
        
        # Delete from vector store - use original user_id format for vector operations
        from backend.rag import VectorStore
        vector_store = VectorStore()
        actual_user_id = "test_user" if user_id == "test_user" else str(user_uuid)
        vector_store.delete_document_embeddings(document_id, actual_user_id)
        
        # Delete from database (chunks will be deleted via cascade)
        db.delete(document)
        db.commit()
        return True
        
    except Exception as e:
        db.rollback()
        print(f"Error deleting document: {e}")
        return False
    finally:
        db.close()